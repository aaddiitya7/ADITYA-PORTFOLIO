"""
ATS-Friendly Resume Generator for Aditya Singh
-----------------------------------------------
Design decisions for ATS compliance:
- Single-column layout (no tables, no text boxes)
- Standard Word Heading styles (H1 for sections, H2 for job titles)
- Arial 11pt body, Arial 12pt bold headers — universal ATS-safe fonts
- 1-inch margins on all sides (standard ATS expectation)
- No hyperlinks — plain text URLs instead
- Horizontal rule dividers between sections
- Skills broken into 3 clear categories with keyword-rich terms
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set page margins to standard 1-inch for ATS parsers."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_horizontal_rule(doc):
    """Add a thin horizontal line between sections (ATS-safe visual separator)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_section_header(doc, text):
    """Add an ATS-parseable section header using Word's built-in Heading 1 style."""
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    # Ensure heading is all-caps and black for ATS + visual consistency
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True
    return h


def add_job_entry(doc, title, company_date, bullets):
    """Add a job entry: title bold, date italic, then bullet points."""
    # Job title + company
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(11)

    # Date on the same paragraph, separated by a tab
    run_date = p.add_run(f'  |  {company_date}')
    run_date.italic = True
    run_date.font.name = 'Arial'
    run_date.font.size = Pt(10)
    run_date.font.color.rgb = RGBColor(80, 80, 80)

    # Bullet points
    for bullet in bullets:
        b = doc.add_paragraph(style='List Bullet')
        b.paragraph_format.space_before = Pt(1)
        b.paragraph_format.space_after = Pt(1)
        run = b.add_run(bullet)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)


def add_skill_line(doc, category, skills_list):
    """Add a skill line: bold category name followed by comma-separated skills."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    cat_run = p.add_run(f'{category}: ')
    cat_run.bold = True
    cat_run.font.name = 'Arial'
    cat_run.font.size = Pt(10.5)
    skill_run = p.add_run(', '.join(skills_list))
    skill_run.font.name = 'Arial'
    skill_run.font.size = Pt(10.5)


def add_cert(doc, cert_name, issuer):
    """Add a single certification as a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    name_run = p.add_run(cert_name)
    name_run.bold = True
    name_run.font.name = 'Arial'
    name_run.font.size = Pt(10.5)
    issuer_run = p.add_run(f'  —  {issuer}')
    issuer_run.font.name = 'Arial'
    issuer_run.font.size = Pt(10.5)
    issuer_run.font.color.rgb = RGBColor(80, 80, 80)


def create_resume():
    doc = Document()

    # ── Global default style ────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

    # ── Margins ─────────────────────────────────────────────────────
    set_margins(doc)

    # ════════════════════════════════════════════════════════════════
    # HEADER
    # ════════════════════════════════════════════════════════════════
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run('ADITYA SINGH')
    name_run.bold = True
    name_run.font.name = 'Arial'
    name_run.font.size = Pt(18)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(2)
    title_run = title_para.add_run('AI-Driven Digital Marketer | SEO | Market Research | E-commerce')
    title_run.bold = True
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(11)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(4)
    contact_run = contact_para.add_run(
        'Kolkata, India  |  aditya818610@gmail.com  |  +91 8335878691  |  '
        'linkedin.com/in/aditya-singh-43162136b'
    )
    contact_run.font.name = 'Arial'
    contact_run.font.size = Pt(10)

    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # PROFESSIONAL SUMMARY
    # ════════════════════════════════════════════════════════════════
    add_section_header(doc, 'PROFESSIONAL SUMMARY')

    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(6)
    summary_run = summary_para.add_run(
        'AI-Driven Digital Marketer and BBA Marketing student at Amity University Kolkata '
        '(Graduating 2027) with 3+ years of experience spanning freelance consulting, e-commerce '
        'operations, and a formal internship. Expert in SEO, Email Marketing, and Market Research, '
        'with a proven ability to leverage AI tools (ChatGPT, Canva AI) and prompt engineering to '
        'rapidly deploy landing pages, web applications, and full marketing funnels. Proficient in '
        'MS Excel, Power BI, Zapier, WordPress, and Google Ads. 8+ industry certifications from '
        'Google, University of Pennsylvania, and Vanderbilt University.'
    )
    summary_run.font.name = 'Arial'
    summary_run.font.size = Pt(10.5)

    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # SKILLS
    # ════════════════════════════════════════════════════════════════
    add_section_header(doc, 'SKILLS')

    add_skill_line(doc, 'Core Marketing',
        ['SEO', 'Digital Marketing', 'Email Marketing', 'E-commerce Strategy',
         'Market Research', 'Consumer Psychology', 'Market Analysis',
         'Google Ads', 'Local Lead Generation', 'Content Optimization'])

    add_skill_line(doc, 'Tools & Software',
        ['MS Excel', 'Power BI', 'Google Search Console', 'Google Trends',
         'Google Ads Manager', 'Zapier', 'WordPress', 'Canva AI',
         'ChatGPT', 'Microsoft PowerPoint'])

    add_skill_line(doc, 'AI & Automation',
        ['Prompt Engineering', 'AI Web Generation', 'Marketing Automation',
         'Zapier Workflows', 'Data Analysis via AI', 'AI-Assisted Content Creation'])

    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # PROFESSIONAL EXPERIENCE
    # ════════════════════════════════════════════════════════════════
    add_section_header(doc, 'PROFESSIONAL EXPERIENCE')

    add_job_entry(doc,
        'Digital Marketing & Market Research Intern  —  Aglo Bond Pvt. Ltd.',
        'June 2025 – August 2025',
        [
            'Conducted in-depth quantitative and qualitative market research (surveys, distributor '
            'interviews, secondary data) to analyse customer behaviour and segment markets across '
            'FMCG categories (pulses, dry fruits, rice).',
            'Used Google Trends and Google Search Console to identify seasonal demand patterns and '
            'content opportunities; findings directly informed campaign messaging.',
            'Compiled and cleaned datasets for lead generation and email campaigns using MS Excel '
            'and Zapier workflows, reducing manual processing time by ~40%.',
            'Supported email and WhatsApp outreach campaigns, maintaining brand consistency across '
            'all touchpoints.',
            'Performed keyword research and supported Google Ads campaign setup; monitored KPIs '
            'and delivered bi-weekly performance reports to stakeholders.',
            'Created Excel dashboards and PowerPoint decks for data storytelling and senior '
            'stakeholder presentations.',
            'Leveraged AI tools (ChatGPT, Canva AI) for research summaries, copy generation, '
            'and visual asset creation.',
        ]
    )

    add_job_entry(doc,
        'E-commerce Seller  —  Amazon India  (Seller ID: A35RU6FGL0700V)',
        'December 2024 – May 2025',
        [
            'Managed a live Amazon India storefront end-to-end: product listings, inventory, '
            'order fulfillment, and customer service.',
            'Conducted keyword research and implemented on-page SEO strategies to improve '
            'organic product visibility and search ranking.',
            'Analysed sales trends and customer feedback to optimise product catalogue and '
            'improve satisfaction scores.',
            'Used data insights to drive consistent revenue growth and improve popularity rank '
            'across multiple product categories.',
        ]
    )

    add_job_entry(doc,
        'Freelance Digital Marketing Consultant  —  Self-Employed',
        'January 2023 – January 2024',
        [
            'Delivered tailored digital marketing services for local businesses: SEO audits, '
            'keyword strategy, and content optimisation.',
            'Executed on-page SEO improvements that improved client organic search traffic.',
            'Managed client communication, project timelines, and deliverables independently.',
        ]
    )

    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # AI-GENERATED MARKETING PROJECTS
    # ════════════════════════════════════════════════════════════════
    add_section_header(doc, 'KEY PROJECTS')

    add_job_entry(doc,
        'A.S. Enterprise B2B Website',
        'Live Deployment — as-enterprise-coral.vercel.app',
        [
            'Architected and deployed a high-converting B2B digital storefront to capture '
            'qualified wholesale inquiries for an export business.',
            'Used AI web generation and prompt engineering to deliver a responsive, '
            'SEO-optimised site with integrated CRM lead capture — deployed in days, '
            'not weeks.',
        ]
    )

    add_job_entry(doc,
        'Forge Task Management Web Application',
        'Live PWA Deployment',
        [
            'Conceptualised and built a Progressive Web App (PWA) from scratch using '
            'AI code generation, demonstrating rapid product development capability.',
            'The app serves as a practical lead magnet and engagement tool, showcasing '
            'full-stack thinking with zero traditional engineering team.',
        ]
    )

    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # CASE STUDIES
    # ════════════════════════════════════════════════════════════════
    add_section_header(doc, 'CASE STUDIES')

    add_job_entry(doc,
        'Diagnosing & Fixing a Silent GA4 Analytics Failure  —  A.S. Enterpriis',
        'March 2026',
        [
            'Identified root cause of 7 days of zero GA4 data: an orphaned Measurement ID '
            'caused by an account-level property migration that silently disconnected '
            "Google's tag-serving backend — undetectable in the GA4 admin UI.",
            'Performed network-level HTTP debugging (diagnosed 404 → 200 status mismatch), '
            're-created the data stream, and deployed the fix via GitHub → Vercel CI/CD '
            'pipeline in under 2 hours, fully restoring data collection.',
        ]
    )

    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # EDUCATION
    # ════════════════════════════════════════════════════════════════
    add_section_header(doc, 'EDUCATION')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run('Amity University Kolkata')
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run('BBA Honours in Marketing  |  2023 – 2027  |  Focus: Marketing Strategy, Branding & Digital Transformation')
    r2.italic = True
    r2.font.name = 'Arial'
    r2.font.size = Pt(10.5)

    add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # CERTIFICATIONS
    # ════════════════════════════════════════════════════════════════
    add_section_header(doc, 'CERTIFICATIONS')

    certs = [
        ('Foundation of Digital Marketing and E-commerce', 'Google / Coursera'),
        ('Attract and Engage Customers with Digital Marketing', 'Google / Coursera'),
        ('Think Outside the Inbox: Email Marketing', 'Google / Coursera'),
        ('From Likes to Leads: Interact with Customers Online', 'Google / Coursera'),
        ('Viral Marketing and How to Craft Contagious Content', 'University of Pennsylvania / Coursera'),
        ('Google SEO Fundamentals', 'UC Davis / Coursera'),
        ('Google Ads for Beginners', 'Coursera'),
        ('AI for Everyone', 'DeepLearning.AI — Andrew Ng'),
        ('Prompt Engineering for ChatGPT', 'Vanderbilt University'),
    ]

    for cert_name, issuer in certs:
        add_cert(doc, cert_name, issuer)

    # ── Save ────────────────────────────────────────────────────────
    output_path = 'Resume/ADITYA_SINGH_RESUME_UPDATED.docx'
    doc.save(output_path)
    print(f'✅  ATS-friendly resume saved → {output_path}')


if __name__ == '__main__':
    create_resume()
